"""
Resume Text Extractor Module
AI Resume Analyzer - Extracts text from PDF and DOCX resume files
"""

import re
from io import BytesIO
from typing import Union

# Import required libraries
try:
    import PyPDF2
except ImportError:
    raise ImportError("PyPDF2 is required. Install it using: pip install PyPDF2")

try:
    import docx
except ImportError:
    raise ImportError("python-docx is required. Install it using: pip install python-docx")

try:
    import spacy
except ImportError:
    raise ImportError("spacy is required. Install it using: pip install spacy")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    raise ImportError("scikit-learn is required. Install it using: pip install scikit-learn")

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    raise ImportError("spaCy model 'en_core_web_sm' not found. Run: python -m spacy download en_core_web_sm")


def extract_text_from_pdf(file: Union[BytesIO, bytes, str]) -> str:
    """
    Extract text from a PDF file object.
    
    Args:
        file: A file object (BytesIO), bytes, or file path to a PDF file
        
    Returns:
        A clean string with all extracted text
        
    Raises:
        ValueError: If the file is corrupted or cannot be read
    """
    try:
        # Handle different input types
        if isinstance(file, str):
            # It's a file path
            with open(file, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
        elif isinstance(file, bytes):
            # It's bytes
            pdf_reader = PyPDF2.PdfReader(BytesIO(file))
        elif hasattr(file, 'read'):
            # It's a file-like object
            file.seek(0)  # Reset to beginning
            pdf_reader = PyPDF2.PdfReader(file)
        else:
            raise ValueError("Invalid file format. Expected file path, bytes, or file object.")
        
        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        
        # Clean up the extracted text
        cleaned_text = clean_text(text)
        
        return cleaned_text
        
    except PyPDF2.errors.PdfReadError as e:
        raise ValueError(f"Error reading PDF file. The file may be corrupted: {str(e)}")
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file: Union[BytesIO, bytes, str]) -> str:
    """
    Extract text from a DOCX file object.
    
    Args:
        file: A file object (BytesIO), bytes, or file path to a DOCX file
        
    Returns:
        A clean string with all extracted text
        
    Raises:
        ValueError: If the file is corrupted or cannot be read
    """
    try:
        # Handle different input types
        if isinstance(file, str):
            # It's a file path
            doc = docx.Document(file)
        elif isinstance(file, bytes):
            # It's bytes
            doc = docx.Document(BytesIO(file))
        elif hasattr(file, 'read'):
            # It's a file-like object
            file.seek(0)  # Reset to beginning
            doc = docx.Document(file)
        else:
            raise ValueError("Invalid file format. Expected file path, bytes, or file object.")
        
        # Extract text from all paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the extracted text
        cleaned_text = clean_text(text)
        
        return cleaned_text
        
    except docx.opc.exceptions.PackageNotFoundError as e:
        raise ValueError(f"Error reading DOCX file. The file may be corrupted: {str(e)}")
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")


def clean_text(text: str) -> str:
    """
    Clean up excessive whitespace and newline characters.
    
    Args:
        text: Raw text string
        
    Returns:
        Cleaned text string
    """
    if not text:
        return ""
    
    # Replace multiple whitespace characters with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove leading/trailing whitespace from each line
    lines = text.split('\n')
    cleaned_lines = [line.strip() for line in lines]
    
    # Join lines, removing empty lines
    cleaned_text = '\n'.join(line for line in cleaned_lines if line)
    
    return cleaned_text.strip()


# Predefined list of common computer science and software engineering skills
SKILLS_LIST = [
    # Programming Languages
    "python", "java", "c++", "c#", "c", "javascript", "typescript", "ruby", "go", "rust",
    "swift", "kotlin", "php", "perl", "scala", "r", "matlab", "objective-c", "shell", "bash",
    
    # Web Development
    "html", "css", "react", "angular", "vue", "node.js", "nodejs", "express", "django", "flask",
    "spring", "asp.net", ".net", "jquery", "bootstrap", "sass", "less", "webpack", "npm",
    "rest", "restful", "graphql", "ajax", "web services",
    
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "oracle", "sqlite", "redis", "cassandra", "elasticsearch",
    "dynamodb", "firebase", "nosql", "database", "data warehousing",
    
    # Data Science & Machine Learning
    "machine learning", "deep learning", "data science", "artificial intelligence", "ai", "ml",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy", "scipy", "matplotlib",
    "nlp", "natural language processing", "computer vision", "neural networks", "statistics",
    "data analysis", "data visualization", "tableau", "power bi",
    
    # Cloud & DevOps
    "aws", "azure", "google cloud", "gcp", "docker", "kubernetes", "jenkins", "terraform",
    "ansible", "puppet", "chef", "ci/cd", "devops", "cloud computing", "microservices",
    "serverless", "lambda",
    
    # Version Control & Project Management
    "git", "github", "gitlab", "bitbucket", "svn", "jira", "confluence", "agile", "scrum",
    "kanban", "waterfall",
    
    # Software Engineering
    "oop", "object oriented", "design patterns", "data structures", "algorithms",
    "software development", "software engineering", "full stack", "backend", "frontend",
    "full-stack", "web development", "mobile development", "ios", "android",
    
    # Testing & Quality
    "unit testing", "integration testing", "test driven development", "tdd", "selenium",
    "junit", "pytest", "automated testing", "qa", "quality assurance",
    
    # Security
    "cybersecurity", "information security", "penetration testing", "owasp", "encryption",
    "security", "network security",
    
    # Other Technical Skills
    "linux", "unix", "windows server", "networking", "tcp/ip", "dns", "http", "https",
    "api", "xml", "json", "yaml", "tomcat", "nginx", "apache", "iis",
    "maven", "gradle", "yarn", "package manager"
]


def extract_skills(text: str) -> list:
    """
    Extract skills from resume text using spaCy phrase matching and string matching.
    
    Args:
        text: Cleaned text string from a resume
        
    Returns:
        A unique list of found skills in lowercase
    """
    if not text:
        return []
    
    # Convert text to lowercase for matching
    text_lower = text.lower()
    
    # Create a spaCy Doc from the text
    doc = nlp(text_lower)
    
    # Set to store unique found skills
    found_skills = set()
    
    # Method 1: Direct string matching with word boundaries
    for skill in SKILLS_LIST:
        # Create pattern with word boundaries to avoid partial matches
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(skill_pattern, text_lower):
            found_skills.add(skill)
    
    # Method 2: spaCy phrase matching for better accuracy
    from spacy.matcher import PhraseMatcher
    
    # Create a new phrase matcher
    phrase_matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
    
    # Add skills as phrase patterns
    phrases = [nlp(skill) for skill in SKILLS_LIST]
    phrase_matcher.add("skills", phrases)
    
    # Find matches in the doc
    matches = phrase_matcher(doc)
    
    for match_id, start, end in matches:
        matched_span = doc[start:end]
        matched_text = matched_span.text.lower()
        # Find the original skill in the list
        for skill in SKILLS_LIST:
            if skill.lower() == matched_text:
                found_skills.add(skill)
                break
    
    # Return unique list sorted alphabetically
    return sorted(list(found_skills))


def calculate_match_score(resume_text: str, job_description_text: str) -> float:
    """
    Calculate the match score between a resume and job description using TF-IDF and cosine similarity.
    
    Args:
        resume_text: The cleaned text from the resume
        job_description_text: The text from the job description
        
    Returns:
        The similarity score as a percentage (0-100), rounded to two decimal places
    """
    # Handle empty or None inputs
    if not resume_text or not job_description_text:
        return 0.0
    
    try:
        # Create TF-IDF vectorizer
        vectorizer = TfidfVectorizer(
            stop_words='english',
            ngram_range=(1, 2),  # Use unigrams and bigrams
            max_features=5000
        )
        
        # Fit and transform both texts
        tfidf_matrix = vectorizer.fit_transform([resume_text, job_description_text])
        
        # Calculate cosine similarity between the two vectors
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Convert to percentage and round to 2 decimal places
        match_percentage = round(similarity * 100, 2)
        
        return match_percentage
        
    except Exception as e:
        raise ValueError(f"Error calculating match score: {str(e)}")


def get_missing_skills(resume_skills: list, job_skills: list) -> list:
    """
    Compare resume skills with job required skills and return missing skills.
    
    Args:
        resume_skills: List of skills found in the resume
        job_skills: List of skills required by the job
        
    Returns:
        List of skills present in job_skills but missing from resume_skills
    """
    if not resume_skills:
        return list(job_skills)
    
    if not job_skills:
        return []
    
    # Convert to sets for efficient comparison (case-insensitive)
    resume_skills_lower = set(skill.lower() for skill in resume_skills)
    job_skills_lower = set(skill.lower() for skill in job_skills)
    
    # Find skills in job_skills that are not in resume_skills
    missing_skills = job_skills_lower - resume_skills_lower
    
    # Return as sorted list
    return sorted(list(missing_skills))


def generate_suggestions(score: float, missing_skills: list) -> list:
    """
    Generate actionable suggestions based on match score and missing skills.
    
    Args:
        score: The match score percentage (0-100)
        missing_skills: List of skills missing from the resume
        
    Returns:
        List of 3 actionable text suggestions
    """
    suggestions = []
    
    # Suggestion 1: Based on overall score
    if score >= 70:
        suggestions.append("Your resume has a strong match! Consider highlighting your relevant experience more prominently.")
    elif score >= 50:
        suggestions.append("Your resume matches moderately. Review the job description and align your skills section better.")
    else:
        suggestions.append("Your resume needs significant improvement. Focus on adding more relevant skills and experience.")
    
    # Suggestion 2: Based on missing skills
    if missing_skills:
        # Show top 3 missing skills
        top_missing = missing_skills[:3]
        skills_text = ", ".join(top_missing)
        if len(missing_skills) > 3:
            skills_text += f" and {len(missing_skills) - 3} more"
        suggestions.append(f"Consider adding these skills to your resume: {skills_text}")
    else:
        suggestions.append("Great job! You have all the key skills required for this position.")
    
    # Suggestion 3: General improvement tip
    if score < 50 and missing_skills:
        suggestions.append("Take online courses or certifications to learn the missing skills mentioned in the job description.")
    elif score >= 50 and score < 70:
        suggestions.append("Tailor your resume's keywords to match the job description more closely for better ATS compatibility.")
    else:
        suggestions.append("Keep your resume updated and continue learning new technologies to stay competitive.")
    
    return suggestions


# Job Role Skill Database - Maps job roles to required skills
JOB_ROLE_SKILLS = {
    "white hat hacker": [
        "python", "penetration testing", "network security", "kali linux", "burp suite",
        "metasploit", "nmap", "sql injection", "xss", "cybersecurity", "information security",
        "owasp", "linux", "encryption", "security", "ethical hacking", "computer forensics",
        "reverse engineering", "malware analysis", "firewalls", "vpn", "tcp/ip"
    ],
    "ethical hacker": [
        "python", "penetration testing", "network security", "kali linux", "burp suite",
        "metasploit", "nmap", "sql injection", "xss", "cybersecurity", "information security",
        "owasp", "linux", "encryption", "security", "ethical hacking", "computer forensics",
        "reverse engineering", "malware analysis", "firewalls", "vpn", "tcp/ip"
    ],
    "web developer": [
        "html", "css", "javascript", "react", "angular", "vue", "node.js", "nodejs",
        "django", "flask", "python", "php", "sql", "mysql", "mongodb", "rest", "restful",
        "graphql", "jquery", "bootstrap", "sass", "less", "webpack", "npm", "git",
        "responsive design", "web services", "api", "json", "xml", "http", "https"
    ],
    "full stack developer": [
        "html", "css", "javascript", "react", "angular", "vue", "node.js", "nodejs",
        "django", "flask", "python", "php", "java", "sql", "mysql", "postgresql", "mongodb",
        "rest", "restful", "graphql", "docker", "git", "aws", "agile", "api", "json"
    ],
    "data scientist": [
        "python", "r", "machine learning", "deep learning", "data science", "tensorflow",
        "pytorch", "keras", "scikit-learn", "pandas", "numpy", "scipy", "matplotlib",
        "data analysis", "data visualization", "statistics", "sql", "tableau", "power bi",
        "natural language processing", "nlp", "computer vision", "neural networks", "ai", "ml"
    ],
    "ai engineer": [
        "python", "machine learning", "deep learning", "artificial intelligence", "ai",
        "tensorflow", "pytorch", "keras", "scikit-learn", "neural networks", "computer vision",
        "natural language processing", "nlp", "data science", "pandas", "numpy", "statistics",
        "docker", "kubernetes", "aws", "azure", "gcp", "ml", "mlops"
    ],
    "software engineer": [
        "python", "java", "c++", "c#", "javascript", "typescript", "git", "docker",
        "kubernetes", "aws", "azure", "agile", "scrum", "oop", "design patterns",
        "data structures", "algorithms", "sql", "rest", "api", "linux", "devops"
    ],
    "devops engineer": [
        "docker", "kubernetes", "jenkins", "terraform", "ansible", "puppet", "chef",
        "aws", "azure", "google cloud", "gcp", "linux", "bash", "shell", "scripting",
        "ci/cd", "devops", "git", "python", "cloud computing", "microservices"
    ],
    "data analyst": [
        "python", "sql", "tableau", "power bi", "data analysis", "data visualization",
        "pandas", "numpy", "statistics", "excel", "r", "machine learning", "ml"
    ],
    "cybersecurity": [
        "cybersecurity", "information security", "network security", "penetration testing",
        "firewalls", "encryption", "owasp", "kali linux", "linux", "python", "siem",
        "vpn", "tcp/ip", "security", "ethical hacking", "computer forensics"
    ]
}


def detect_job_role(job_description: str) -> tuple:
    """
    Detect the job role from the job description text.
    
    Args:
        job_description: The text from the job description
        
    Returns:
        A tuple of (detected_role, confidence_score) or (None, 0) if no role detected
    """
    if not job_description:
        return None, 0
    
    job_desc_lower = job_description.lower()
    role_matches = {}
    
    # Define role keywords to search for (more flexible)
    role_keywords = {
        "white hat hacker": ["white hat", "ethical hacker", "ethical hacking", "penetration tester"],
        "ethical hacker": ["ethical hacker", "ethical hacking", "penetration tester"],
        "web developer": ["web developer", "web development", "frontend developer", "front-end developer"],
        "full stack developer": ["full stack", "fullstack", "full-stack", "full stack developer"],
        "data scientist": ["data scientist", "data science"],
        "ai engineer": ["ai engineer", "ai developer", "artificial intelligence", "ml engineer", "machine learning engineer"],
        "software engineer": ["software engineer", "software developer"],
        "devops engineer": ["devops", "devops engineer", "sre", "site reliability"],
        "data analyst": ["data analyst", "data analysis", "analytics"],
        "cybersecurity": ["cybersecurity", "cyber security", "security analyst", "infosec"]
    }
    
    # Check each role against the job description
    for role, keywords in role_keywords.items():
        match_count = 0
        for keyword in keywords:
            if keyword in job_desc_lower:
                match_count += 1
        
        if match_count > 0:
            role_matches[role] = match_count
    
    if not role_matches:
        return None, 0
    
    # Return the role with highest match
    detected_role = max(role_matches, key=role_matches.get)
    confidence = min(role_matches[detected_role] * 25, 100)  # Scale to 100%
    
    return detected_role, confidence


def get_role_required_skills(role: str) -> list:
    """
    Get the required skills for a specific job role.
    
    Args:
        role: The job role name
        
    Returns:
        List of required skills for the role (lowercase)
    """
    if not role:
        return []
    
    role_lower = role.lower()
    
    # Try to find exact match first
    if role_lower in JOB_ROLE_SKILLS:
        return JOB_ROLE_SKILLS[role_lower]
    
    # Try partial match
    for job_role in JOB_ROLE_SKILLS.keys():
        if role_lower in job_role or job_role in role_lower:
            return JOB_ROLE_SKILLS[job_role]
    
    return []


def calculate_role_based_score(resume_skills: list, required_skills: list) -> float:
    """
    Calculate the match score based on role-specific skills.
    
    Formula: score = (matched_skills / required_skills) * 100
    
    Args:
        resume_skills: List of skills found in the resume
        required_skills: List of required skills for the job role
        
    Returns:
        Score as a percentage (0-100), rounded to 2 decimal places
    """
    if not required_skills:
        return 0.0
    
    if not resume_skills:
        return 0.0
    
    # Convert to sets for comparison (case-insensitive)
    resume_skills_lower = set(skill.lower() for skill in resume_skills)
    required_skills_lower = set(skill.lower() for skill in required_skills)
    
    # Find matched skills
    matched_skills = resume_skills_lower.intersection(required_skills_lower)
    
    # Calculate score
    score = (len(matched_skills) / len(required_skills_lower)) * 100
    
    return round(score, 2)


def get_matched_skills_for_role(resume_skills: list, required_skills: list) -> list:
    """
    Get the list of skills that match between resume and role requirements.
    
    Args:
        resume_skills: List of skills found in the resume
        required_skills: List of required skills for the job role
        
    Returns:
        List of matched skills (lowercase)
    """
    if not resume_skills or not required_skills:
        return []
    
    resume_skills_lower = set(skill.lower() for skill in resume_skills)
    required_skills_lower = set(skill.lower() for skill in required_skills)
    
    matched = resume_skills_lower.intersection(required_skills_lower)
    
    return sorted(list(matched))


def get_role_missing_skills(resume_skills: list, required_skills: list) -> list:
    """
    Get the missing skills required for the role but not in the resume.
    
    Args:
        resume_skills: List of skills found in the resume
        required_skills: List of required skills for the job role
        
    Returns:
        List of missing skills (lowercase)
    """
    if not required_skills:
        return []
    
    if not resume_skills:
        return sorted([skill.lower() for skill in required_skills])
    
    resume_skills_lower = set(skill.lower() for skill in resume_skills)
    required_skills_lower = set(skill.lower() for skill in required_skills)
    
    missing = required_skills_lower - resume_skills_lower
    
    return sorted(list(missing))


def generate_role_recommendations(missing_skills: list, role: str) -> str:
    """
    Generate learning recommendations based on missing skills for a specific role.
    
    Args:
        missing_skills: List of skills missing from the resume
        role: The job role name
        
    Returns:
        A recommendation string
    """
    if not missing_skills:
        return f"Congratulations! You have all the required skills for the {role} position!"
    
    # Get top 3 most important missing skills to highlight
    top_skills = missing_skills[:3]
    
    skills_text = ", ".join([skill.title() for skill in top_skills])
    
    if len(missing_skills) > 3:
        recommendation = f"You should learn {skills_text}, and {len(missing_skills) - 3} more skills to qualify for this {role} role."
    else:
        recommendation = f"You should learn {skills_text} to qualify for this {role} role."
    
    return recommendation


# Example usage and testing
if __name__ == "__main__":
    # Test with sample files (uncomment and modify paths as needed)
    
    # Test PDF extraction
    # try:
    #     with open("sample_resume.pdf", "rb") as pdf_file:
    #         pdf_text = extract_text_from_pdf(pdf_file)
    #         print("PDF Text extracted successfully!")
    #         print(f"Length: {len(pdf_text)} characters")
    # except ValueError as e:
    #     print(f"PDF Error: {e}")
    
    # Test DOCX extraction
    # try:
    #     with open("sample_resume.docx", "rb") as docx_file:
    #         docx_text = extract_text_from_docx(docx_file)
    #         print("DOCX Text extracted successfully!")
    #         print(f"Length: {len(docx_text)} characters")
    # except ValueError as e:
    #     print(f"DOCX Error: {e}")
    
    print("Resume Extractor Module loaded successfully!")
    print("Functions available:")
    print("  - extract_text_from_pdf(file)")
    print("  - extract_text_from_docx(file)")
    print("  - clean_text(text)")
    print("  - extract_skills(text)")
    print("  - calculate_match_score(resume_text, job_description_text)")
    print("  - get_missing_skills(resume_skills, job_skills)")
    print("  - generate_suggestions(score, missing_skills)")
    print("  - detect_job_role(job_description)")
    print("  - get_role_required_skills(role)")
    print("  - calculate_role_based_score(resume_skills, required_skills)")
    print("  - generate_role_recommendations(missing_skills, role)")

